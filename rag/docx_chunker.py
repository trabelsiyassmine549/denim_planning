"""
rag/docx_chunker.py — Extracts and chunks the RAG knowledge base DOCX
=======================================================================
Usage (run once, or at startup if chunks.json is missing):

    python -m rag.docx_chunker

Produces:  rag/chunks.json   — list of {"text": "...", "section": "..."} objects

The rag_engine.py then loads this file instead of the hardcoded DOMAIN_CHUNKS list.

Strategy
--------
- Parse the DOCX section by section (heading → body text)
- Each top-level section (H1/H2) becomes one or more chunks
- Code blocks are kept together (they are small)
- Tables are converted to flat text rows
- Max chunk size : 1400 characters (raised from 1200 to accommodate dense rule sections
  like PARTIE IX without breaking rules mid-sentence across chunk boundaries)
- Min chunk size :  80 characters  (ignore near-empty paragraphs)

Staleness detection
-------------------
load_chunks() compares the DOCX mtime against chunks.json mtime.
If the DOCX is newer, chunks.json is automatically rebuilt so stale
chunks are never served after a knowledge base update.

Split strategy
--------------
_split_text() tries boundaries in order:
  1. Newline (preserves bullet/rule structure)
  2. Sentence boundary ". "
  3. Hard cut at MAX_CHUNK (last resort)
"""

import json
import os
import re
from pathlib import Path

try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
except ImportError:
    raise ImportError("python-docx is required: pip install python-docx --break-system-packages")

DOCX_PATH   = Path(__file__).parent / "rag_knowledge_base_v2.docx"
OUTPUT_PATH = Path(__file__).parent / "chunks.json"
MAX_CHUNK   = 1400   # characters — raised from 1200 to keep dense rule sections intact
MIN_CHUNK   = 80


# ── Helpers ─────────────────────────────────────────────────────────────────

def _heading_level(para) -> int:
    """Return 1/2/3 for heading paragraphs, 0 otherwise."""
    style = para.style.name if para.style else ""
    if style.startswith("Heading 1"):
        return 1
    if style.startswith("Heading 2"):
        return 2
    if style.startswith("Heading 3"):
        return 3
    return 0


def _is_code(para) -> bool:
    """Heuristic: paragraph with Courier New font = code block."""
    for run in para.runs:
        font = run.font.name or ""
        if "courier" in font.lower():
            return True
    return False


def _table_to_text(table) -> str:
    """Convert a DOCX table to readable flat text."""
    lines = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        line  = " | ".join(cells)
        if i == 0:
            lines.append(line)
            lines.append("-" * min(len(line), 80))
        else:
            lines.append(line)
    return "\n".join(lines)


def _split_text(text: str, section: str, max_size: int = MAX_CHUNK):
    """
    Split a long text into chunks of at most max_size characters.

    Split boundary priority (to preserve rule and bullet structure):
      1. Newline — keeps list items and rules together
      2. Sentence boundary ". " — fallback for prose paragraphs
      3. Hard cut at max_size — last resort only
    """
    chunks = []
    while len(text) > max_size:
        # 1. Prefer splitting at a newline boundary
        cut = text.rfind("\n", 0, max_size)
        if cut > max_size // 2:          # only use if it's not too far back
            cut += 1                     # include the newline
        else:
            # 2. Fall back to sentence boundary
            cut = text.rfind(". ", 0, max_size)
            if cut != -1:
                cut += 1                 # include the period
            else:
                # 3. Hard cut — last resort
                cut = max_size
        chunks.append({"text": text[:cut].strip(), "section": section})
        text = text[cut:].strip()
    if text and len(text) >= MIN_CHUNK:
        chunks.append({"text": text, "section": section})
    return chunks


# ── Main extractor ───────────────────────────────────────────────────────────

def extract_chunks(docx_path: Path) -> list:
    """
    Parse the DOCX and return a list of chunk dicts:
    {"text": "...", "section": "<H1 > H2 title>"}
    """
    doc          = DocxDocument(docx_path)
    chunks       = []
    current_h1   = ""
    current_h2   = ""
    buffer_text  = ""
    buffer_label = ""

    def flush(label=None):
        nonlocal buffer_text, buffer_label
        t = buffer_text.strip()
        l = label or buffer_label
        if t and len(t) >= MIN_CHUNK:
            chunks.extend(_split_text(t, l))
        buffer_text  = ""
        buffer_label = l or buffer_label

    # Iterate body elements in order (paragraphs + tables)
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "tbl":
            # Table — convert to text and add to buffer
            # Find the Table object matching this XML element
            for tbl in doc.tables:
                if tbl._tbl is child:
                    buffer_text += "\n" + _table_to_text(tbl) + "\n"
                    break

        elif tag == "p":
            # Find the Paragraph object
            para = None
            for p in doc.paragraphs:
                if p._p is child:
                    para = p
                    break
            if para is None:
                continue

            lvl  = _heading_level(para)
            text = para.text.strip()

            if lvl == 1:
                flush()
                current_h1   = text
                current_h2   = ""
                buffer_label = current_h1
                # Add the heading itself as the first line of the next buffer
                buffer_text  = f"=== {text} ===\n"

            elif lvl == 2:
                flush()
                current_h2   = text
                buffer_label = f"{current_h1} > {current_h2}"
                buffer_text  = f"--- {text} ---\n"

            elif lvl == 3:
                # Sub-heading — keep in buffer but note it
                if text:
                    buffer_text += f"\n[{text}]\n"

            else:
                if not text:
                    continue
                is_code = _is_code(para)
                if is_code:
                    buffer_text += text + "\n"
                else:
                    # Flush if adding this paragraph would exceed max chunk
                    if len(buffer_text) + len(text) + 2 > MAX_CHUNK:
                        flush()
                    buffer_text += text + "\n"

    flush()  # flush whatever remains

    # Post-process: deduplicate, remove near-duplicates and boilerplate
    seen = set()
    final = []
    for c in chunks:
        # FIX ISSUE 9: preserve internal newlines (bullet/rule structure).
        # Only collapse 3+ consecutive blank lines to 2 — do NOT flatten all whitespace.
        t = re.sub(r'\n{3,}', '\n\n', c["text"]).strip()
        if not t or len(t) < MIN_CHUNK:
            continue
        key = t[:120]
        if key in seen:
            continue
        seen.add(key)
        c["text"] = t
        final.append(c)

    return final


# ── Entry point ──────────────────────────────────────────────────────────────

def build_chunks():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(
            f"DOCX not found: {DOCX_PATH}\n"
            "Place rag_knowledge_base_v2.docx in the rag/ folder."
        )
    print(f"[CHUNKER] Parsing {DOCX_PATH} ...")
    chunks = extract_chunks(DOCX_PATH)
    OUTPUT_PATH.write_text(json.dumps(chunks, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[CHUNKER] {len(chunks)} chunks written to {OUTPUT_PATH}")
    return chunks


def load_chunks() -> list:
    """
    Load chunks from the JSON cache, auto-rebuilding if stale or missing.

    Staleness rule: if the DOCX file is newer than chunks.json, the cache
    is considered stale and is rebuilt automatically. This means updating
    rag_knowledge_base_v2.docx and restarting FastAPI is sufficient —
    no manual chunker run needed.

    Returns a list of dicts: [{"text": "...", "section": "..."}, ...]
    Callers that need plain strings should do: [c["text"] for c in load_chunks()]
    """
    needs_rebuild = False

    if not OUTPUT_PATH.exists():
        print("[CHUNKER] chunks.json not found — building from DOCX...")
        needs_rebuild = True
    elif DOCX_PATH.exists():
        docx_mtime   = DOCX_PATH.stat().st_mtime
        chunks_mtime = OUTPUT_PATH.stat().st_mtime
        if docx_mtime > chunks_mtime:
            print(
                f"[CHUNKER] rag_knowledge_base_v2.docx is newer than chunks.json "
                f"({docx_mtime:.0f} > {chunks_mtime:.0f}) — rebuilding..."
            )
            needs_rebuild = True

    if needs_rebuild:
        build_chunks()

    raw = json.loads(OUTPUT_PATH.read_text(encoding="utf-8"))
    chunks = [c for c in raw if c.get("text")]
    print(f"[CHUNKER] Loaded {len(chunks)} chunks from {OUTPUT_PATH}")
    return chunks


if __name__ == "__main__":
    build_chunks()